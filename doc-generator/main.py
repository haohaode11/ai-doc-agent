from fastapi import FastAPI
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import os
import re

app = FastAPI()

OUTPUT_DIR = "generated"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ================= 工具函数 =================

def safe_filename(name: str) -> str:
    name = name.strip()
    name = re.sub(r'[\\/:*?"<>|]', "_", name)
    return name[:100] if name else "document"


def is_table_line(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_separator_row(row):
    return all(re.fullmatch(r"-+", cell) for cell in row)


def parse_markdown_table(lines):
    rows = []
    for line in lines:
        row = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(row)
    return [r for r in rows if not is_separator_row(r)]


# ================= 请求模型 =================

class GenerateRequest(BaseModel):
    title: str
    content: str
    file_type: str = "word"   # word | excel


# ================= Word 生成 =================

def generate_word(title: str, content: str, filepath: str):
    doc = Document()
    doc.add_heading(title, level=1)

    lines = content.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # ---- Markdown 表格 ----
        if is_table_line(line):
            table_lines = []
            while i < len(lines) and is_table_line(lines[i]):
                table_lines.append(lines[i])
                i += 1

            table_data = parse_markdown_table(table_lines)

            if table_data:
                rows, cols = len(table_data), len(table_data[0])
                table = doc.add_table(rows=rows, cols=cols)
                table.style = "Table Grid"

                for r in range(rows):
                    for c in range(cols):
                        table.rows[r].cells[c].text = table_data[r][c]

        # ---- 普通文本 ----
        else:
            doc.add_paragraph(line)
            i += 1

    doc.save(filepath)


# ================= Excel 生成（多 Sheet + 美化） =================

def generate_excel(title: str, content: str, filepath: str):
    wb = Workbook()
    wb.remove(wb.active)

    lines = content.splitlines()
    i = 0
    sheet_index = 1

    thin = Side(style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    while i < len(lines):
        line = lines[i]

        if is_table_line(line):
            table_lines = []
            while i < len(lines) and is_table_line(lines[i]):
                table_lines.append(lines[i])
                i += 1

            table_data = parse_markdown_table(table_lines)
            if not table_data:
                continue

            sheet_name = f"Sheet{sheet_index}"
            sheet_index += 1
            ws = wb.create_sheet(title=sheet_name)

            # 写入数据
            for r_idx, row in enumerate(table_data, start=1):
                for c_idx, value in enumerate(row, start=1):
                    cell = ws.cell(row=r_idx, column=c_idx, value=value)
                    cell.border = border

                    # 表头样式
                    if r_idx == 1:
                        cell.font = Font(bold=True)
                        cell.alignment = Alignment(horizontal="center")

            # 列宽自适应
            for col_idx in range(1, len(table_data[0]) + 1):
                col_letter = get_column_letter(col_idx)
                max_len = max(
                    len(str(ws.cell(row=r, column=col_idx).value or ""))
                    for r in range(1, len(table_data) + 1)
                )
                ws.column_dimensions[col_letter].width = min(max_len + 4, 30)

        else:
            i += 1

    wb.save(filepath)


# ================= 主接口 =================

@app.post("/generate")
def generate(req: GenerateRequest):
    filename = f"{safe_filename(req.title)}"
    filepath = ""

    if req.file_type == "word":
        filename += ".docx"
        filepath = os.path.join(OUTPUT_DIR, filename)
        generate_word(req.title, req.content, filepath)

    elif req.file_type == "excel":
        filename += ".xlsx"
        filepath = os.path.join(OUTPUT_DIR, filename)
        generate_excel(req.title, req.content, filepath)

    else:
        return {"error": "unsupported file_type"}

    return {
        "file_url": f"http://localhost:5001/download/{filename}"
    }


# ================= 下载接口 =================

@app.get("/download/{filename}")
def download(filename: str):
    path = os.path.join(OUTPUT_DIR, filename)
    return FileResponse(path, filename=filename)
